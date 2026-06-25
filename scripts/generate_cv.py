"""Generate polished DevOps CV — English & Hebrew (ATS-friendly, no keyword spam)."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "cv"

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
TEAL = RGBColor(0x0E, 0x74, 0x90)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

GITHUB = "github.com/EstiGenauer"
PROJECT = "github.com/EstiGenauer/ai-learning-platform"


@dataclass
class SkillGroup:
    title: str
    items: str


@dataclass
class CvContent:
    filename_stem: str
    rtl: bool
    name: str
    subtitle: str
    phone: str
    email: str
    github: str
    project: str
    location: str
    summary: str
    skill_groups: list[SkillGroup]
    project_title: str
    project_application_label: str
    project_application: str
    project_devops_label: str
    project_devops_bullets: list[str]
    education: list[str]
    experience_title: str
    experience_body: str
    languages: str
    soft_skills: str


EN = CvContent(
    filename_stem="Esther_Genauer_CV_EN_DevOps",
    rtl=False,
    name="Esther Genauer",
    subtitle="Junior DevOps Engineer",
    phone="052-7627610",
    email="genauer1997@gmail.com",
    github=GITHUB,
    project=PROJECT,
    location="Israel",
    summary=(
        "Junior DevOps Engineer with hands-on experience building a cloud-native portfolio platform "
        "end-to-end. Skilled in CI/CD automation (GitHub Actions), containerization (Docker, Kubernetes, "
        "Helm), Infrastructure as Code (Terraform, AWS EKS), and observability (Prometheus, Grafana, Loki). "
        "Formally trained in DevOps systems management and .NET development (C#, WPF, WCF). Delivered a "
        "production-style system with automated testing, image publishing, and infrastructure documentation — "
        "using AI-assisted workflows (Cursor, Gemini) while maintaining full architectural ownership."
    ),
    skill_groups=[
        SkillGroup("Cloud & Infrastructure", "AWS (EKS) · Terraform · Kubernetes · Helm · Docker · Linux · Git"),
        SkillGroup("CI/CD & Automation", "GitHub Actions · Azure DevOps · Docker Compose · Docker Hub · GHCR"),
        SkillGroup("Observability", "Prometheus · Grafana · Loki · Promtail · Serilog · /metrics"),
        SkillGroup("Formal Training", "C# · .NET · WPF · WCF · SQL · Python · Java · C"),
        SkillGroup("Portfolio Delivery", "PostgreSQL · REST APIs · React · TypeScript · OpenAI (AI-assisted)"),
        SkillGroup("AI Tools", "Cursor · Google Gemini · GPT — IaC, pipelines, debugging, docs"),
    ],
    project_title="AI Learning Platform — Microservices DevOps System (Flagship Project)",
    project_application_label="Application",
    project_application=(
        "Developed a cloud-native microservices platform: a React frontend and three independent .NET services "
        "(authentication, catalog, and AI/OpenAI GPT-4o) behind an NGINX API gateway, each owning its own PostgreSQL "
        "database, with REST service-to-service communication."
    ),
    project_devops_label="DevOps & Infrastructure",
    project_devops_bullets=[
        "Architected a microservices system (3 .NET services, database-per-service) with an NGINX API gateway and REST inter-service calls",
        "Containerized every service with Docker & Docker Compose",
        "Authored Kubernetes manifests for each service (Deployments, Services, gateway, NGINX Ingress)",
        "Built CI/CD pipelines with GitHub Actions (matrix build & test, Docker image publish to Docker Hub & GHCR)",
        "Implemented Infrastructure as Code with Terraform (local Kubernetes + AWS EKS)",
        "Configured monitoring & logging with Prometheus, Grafana, Loki, Promtail, and /metrics instrumentation",
        "Delivered automated tests plus a full-stack Docker Compose integration smoke test with a mock AI service",
        "Leveraged AI-assisted development tools (Cursor, Gemini) to accelerate implementation while independently "
        "owning system architecture, infrastructure, CI/CD pipelines, deployment, testing, monitoring, and operations",
    ],
    education=[
        "Mercaz Beit Yaakov Seminary (2024–2026) — DevOps & Technology Systems Management · .NET (C#, WPF, WCF)",
        "Pninei Chein High School (2020–2024) — Full Bagrut with Honors · Computer Science (11 units) · Math 5 · English 5",
    ],
    experience_title="Private Tutor — Mathematics & Computer Science",
    experience_body=(
        "Taught high school students (4–5 unit level). Developed analytical thinking, problem decomposition, "
        "and clear technical communication — skills directly applicable to troubleshooting and documentation in DevOps."
    ),
    languages="Hebrew (Native)  ·  English (Native)",
    soft_skills=(
        "Systems thinking · End-to-end ownership · Fast independent learning · "
        "Technical communication · Documentation discipline · DevOps mindset"
    ),
)

HE = CvContent(
    filename_stem="Esther_Genauer_CV_HE_DevOps",
    rtl=True,
    name="אסתר גנאור",
    subtitle="Junior DevOps Engineer",
    phone="052-7627610",
    email="genauer1997@gmail.com",
    github=GITHUB,
    project=PROJECT,
    location="ישראל",
    summary=(
        "Junior DevOps Engineer עם ניסיון hands-on בבניית פלטפורמת portfolio cloud-native מקצה לקצה. "
        "מיומנות ב-CI/CD (GitHub Actions), containerization (Docker, Kubernetes, Helm), "
        "Infrastructure as Code (Terraform, AWS EKS) ו-observability (Prometheus, Grafana, Loki). "
        "הכשרה פורמלית ב-DevOps וב-.NET (C#, WPF, WCF). מערכת production-style עם בדיקות אוטומטיות, "
        "פרסום images ותיעוד תשתית — עם AI-assisted workflow (Cursor, Gemini) ואחריות ארכיטקטונית מלאה."
    ),
    skill_groups=[
        SkillGroup("Cloud & Infrastructure", "AWS (EKS) · Terraform · Kubernetes · Helm · Docker · Linux · Git"),
        SkillGroup("CI/CD & Automation", "GitHub Actions · Azure DevOps · Docker Compose · Docker Hub · GHCR"),
        SkillGroup("Observability", "Prometheus · Grafana · Loki · Promtail · Serilog · /metrics"),
        SkillGroup("הכשרה פורמלית", "C# · .NET · WPF · WCF · SQL · Python · Java · C"),
        SkillGroup("Portfolio", "PostgreSQL · REST · React · TypeScript · OpenAI (AI-assisted)"),
        SkillGroup("כלי AI", "Cursor · Gemini · GPT — IaC, pipelines, debugging, docs"),
    ],
    project_title="AI Learning Platform — מערכת Microservices DevOps (פרויקט מרכזי)",
    project_application_label="אפליקציה",
    project_application=(
        "פיתוח פלטפורמת microservices מסוג cloud-native: React frontend ושלושה שירותי .NET עצמאיים "
        "(authentication, catalog, ו-AI/OpenAI GPT-4o) מאחורי NGINX API gateway, כשכל שירות בעל database "
        "(PostgreSQL) משלו ותקשורת REST בין השירותים."
    ),
    project_devops_label="DevOps ותשתיות",
    project_devops_bullets=[
        "ארכיטקטורת microservices (3 שירותי .NET, database-per-service) עם NGINX API gateway ותקשורת REST בין שירותים",
        "Containerization של כל שירות עם Docker ו-Docker Compose",
        "כתיבת Kubernetes manifests לכל שירות (Deployments, Services, gateway, NGINX Ingress)",
        "CI/CD ב-GitHub Actions (matrix build & test, פרסום images ל-Docker Hub & GHCR)",
        "Infrastructure as Code ב-Terraform (K8s מקומי + AWS EKS)",
        "Monitoring & logging: Prometheus, Grafana, Loki, Promtail, /metrics instrumentation",
        "בדיקות אוטומטיות + full-stack Docker Compose integration smoke test עם mock AI service",
        "שימוש ב-AI-assisted tools (Cursor, Gemini) להאצת פיתוח, תוך אחריות עצמאית על ארכיטקטורת מערכת, "
        "תשתית, CI/CD pipelines, פריסה, בדיקות, monitoring ותפעול",
    ],
    education=[
        "מרכז בית יעקב (2024–2026) — DevOps & Technology Systems Management · .NET (C#, WPF, WCF)",
        "פניני חן תיכון (2020–2024) — בגרות מלאה בהצטיינות · מחשב 5 יח״ל · מתמטיקה 5 · אנגלית 5",
    ],
    experience_title="מורה פרטית — מתמטיקה ומחשב",
    experience_body=(
        "ליווי תלמידות תיכון (4–5 יח״ל) במתמטיקה ומחשב. חשיבה אנליטית, פירוק בעיות "
        "והסבר טכני ברור — רלוונטי ל-troubleshooting ו-documentation ב-DevOps."
    ),
    languages="עברית (שפת אם)  ·  אנגלית (שפת אם)",
    soft_skills=(
        "חשיבה מערכתית · אחריות end-to-end · למידה עצמית מהירה · "
        "תקשורת טכנית · תיעוד · DevOps mindset"
    ),
)


def shade_cell(cell, fill: str) -> None:
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))
    set_cell_margins(cell)


def set_cell_margins(cell, top=40, bottom=40, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def styled_run(p, text, *, size=9, bold=False, color=NAVY, italic=False):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def set_rtl(p):
    p._element.get_or_add_pPr().append(OxmlElement("w:bidi"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_header_banner(doc: Document, c: CvContent) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, "1E3A5F")

    align = WD_ALIGN_PARAGRAPH.RIGHT if c.rtl else WD_ALIGN_PARAGRAPH.CENTER

    p_name = cell.paragraphs[0]
    p_name.alignment = align
    if c.rtl:
        set_rtl(p_name)
    styled_run(p_name, c.name, size=22, bold=True, color=WHITE)

    p_sub = cell.add_paragraph()
    p_sub.alignment = align
    if c.rtl:
        set_rtl(p_sub)
    styled_run(p_sub, c.subtitle, size=11, color=RGBColor(0xBF, 0xDB, 0xFE))

    p_contact = cell.add_paragraph()
    p_contact.alignment = align
    if c.rtl:
        set_rtl(p_contact)
    styled_run(
        p_contact,
        f"{c.phone}  ·  {c.email}  ·  {c.location}  ·  GitHub: {c.github}",
        size=8.5,
        color=RGBColor(0xE2, 0xE8, 0xF0),
    )

    p_link = cell.add_paragraph()
    p_link.alignment = align
    if c.rtl:
        set_rtl(p_link)
    styled_run(p_link, f"Portfolio: {c.project}", size=8.5, color=RGBColor(0x7D, 0xD3, 0xFC))


def add_section_title(doc: Document, title: str, rtl_flag: bool) -> None:
    p = doc.add_paragraph()
    if rtl_flag:
        set_rtl(p)
    styled_run(p, title, size=11, bold=True, color=TEAL)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0E7490")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_body(doc: Document, text: str, rtl_flag: bool, size=9.5) -> None:
    p = doc.add_paragraph()
    if rtl_flag:
        set_rtl(p)
    styled_run(p, text, size=size, color=GRAY)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)


def add_skills_grid(doc: Document, groups: list[SkillGroup], rtl_flag: bool) -> None:
    table = doc.add_table(rows=0, cols=2)
    for i in range(0, len(groups), 2):
        row = table.add_row()
        for col in range(2):
            idx = i + col
            cell = row.cells[col]
            if idx >= len(groups):
                continue
            g = groups[idx]
            shade_cell(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if rtl_flag:
                set_rtl(p)
            styled_run(p, g.title, size=8.5, bold=True, color=NAVY)
            p2 = cell.add_paragraph()
            if rtl_flag:
                set_rtl(p2)
            styled_run(p2, g.items, size=8.5, color=GRAY)


def add_bullets(doc: Document, items: list[str], rtl_flag: bool) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if rtl_flag:
            set_rtl(p)
        styled_run(p, item, size=9, color=GRAY)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1


def build_doc(c: CvContent) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    add_header_banner(doc, c)

    labels = (
        ("Professional Summary", "Technical Skills", "Featured Project", "Education", "Experience", "Languages & Strengths")
        if not c.rtl
        else ("סיכום מקצועי", "מיומנויות טכניות", "פרויקט מרכזי", "השכלה", "ניסיון", "שפות וחוזקות")
    )

    add_section_title(doc, labels[0], c.rtl)
    add_body(doc, c.summary, c.rtl)

    add_section_title(doc, labels[1], c.rtl)
    add_skills_grid(doc, c.skill_groups, c.rtl)

    add_section_title(doc, labels[2], c.rtl)
    p = doc.add_paragraph()
    if c.rtl:
        set_rtl(p)
    styled_run(p, c.project_title, size=10, bold=True, color=NAVY)
    add_body(doc, f"https://{c.project}", c.rtl, size=8.5)
    ap = doc.add_paragraph()
    if c.rtl:
        set_rtl(ap)
    styled_run(ap, f"{c.project_application_label}:", size=9, bold=True, color=NAVY)
    add_body(doc, c.project_application, c.rtl, size=9)
    dp = doc.add_paragraph()
    if c.rtl:
        set_rtl(dp)
    styled_run(dp, f"{c.project_devops_label}:", size=9, bold=True, color=NAVY)
    add_bullets(doc, c.project_devops_bullets, c.rtl)

    add_section_title(doc, labels[3], c.rtl)
    for line in c.education:
        bp = doc.add_paragraph(style="List Bullet")
        if c.rtl:
            set_rtl(bp)
        styled_run(bp, line, size=9, color=GRAY)

    add_section_title(doc, labels[4], c.rtl)
    ep = doc.add_paragraph()
    if c.rtl:
        set_rtl(ep)
    styled_run(ep, c.experience_title, size=9.5, bold=True, color=NAVY)
    add_body(doc, c.experience_body, c.rtl, size=9)

    add_section_title(doc, labels[5], c.rtl)
    add_body(doc, c.languages, c.rtl, size=9)
    add_body(doc, c.soft_skills, c.rtl, size=9)

    return doc


def save_pair(c: CvContent) -> None:
    docx = OUTPUT_DIR / f"{c.filename_stem}.docx"
    pdf = OUTPUT_DIR / f"{c.filename_stem}.pdf"
    build_doc(c).save(docx)
    print(f"Created: {docx}")
    try:
        from docx2pdf import convert
        convert(str(docx), str(pdf))
        print(f"Created: {pdf}")
    except Exception as exc:
        print(f"PDF skipped ({exc})")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    save_pair(EN)
    save_pair(HE)
    print("Done.")


if __name__ == "__main__":
    main()
